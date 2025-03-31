import inquirer
import os
from docx import Document
from docx.shared import Inches
import re


# Function to list only markdown files in the current directory, excluding paper.md and CONTRIBUTING.md
def get_files_in_current_directory():
    excluded_files = {"paper.md", "CONTRIBUTING.md"}
    return [
        f
        for f in os.listdir(".")
        if os.path.isfile(f) and f.endswith(".md") and f not in excluded_files
    ]


# Prompt user to select a file
files = get_files_in_current_directory()
questions = [
    inquirer.List("file", message="Select the file to process", choices=files),
]
selected_file = inquirer.prompt(questions)["file"]

# Load the content of the selected file
with open(selected_file, "r", encoding="utf-8") as file:
    lines = file.readlines()


def parse_comments(lines):
    comments = []
    current_id = None
    current_comment = []
    current_response = []
    mode = "id"  # Tracks if we are in ID, comment, or response mode
    in_multiline_comment = False

    for line in lines:
        line = line.strip()

        # Check for start of a multiline comment
        if line.startswith("<!--"):
            in_multiline_comment = True

        # Check for end of a multiline comment
        if line.endswith("-->"):
            in_multiline_comment = False
            continue

        # Skip lines that are within a multiline comment
        if in_multiline_comment:
            continue

        # Skip single line comments
        if line.startswith("<!--") and line.endswith("-->"):
            continue

        # Start of a new comment section
        if line.startswith("# "):
            # Save the previous comment if we were working on one
            if current_id is not None:
                comments.append(
                    (
                        current_id,
                        "\n".join(current_comment),
                        "\n".join(current_response),
                    )
                )
                current_comment, current_response = [], []  # Reset for the new section

            # Set the new comment ID and switch to comment mode
            current_id = line[2:].strip()
            mode = "comment"

        # Start of the response section
        elif line.startswith("> "):
            mode = "response"
            current_response.append(
                line[2:].strip()
            )  # Start with the first line of response text

        # Append lines to the current section (comment or response)
        else:
            if mode == "comment":
                current_comment.append(line)
            elif mode == "response":
                current_response.append(line)

    # Append the last collected comment and response after the loop
    if current_id is not None:
        comments.append(
            (current_id, "\n".join(current_comment), "\n".join(current_response))
        )

    return comments


def add_markdown_text(paragraph, text):
    """Adds text with Markdown-style bold and italic formatting to a Word paragraph."""
    # Patterns for bold (**text** or __text__) and italic (*text* or _text_)
    bold_italic_pattern = re.compile(r"(\*\*\*)(.*?)\1")  # Matches ***bold italic***
    bold_pattern = re.compile(r"(\*\*|__)(.*?)\1")  # Matches **bold** or __bold__
    italic_pattern = re.compile(r"(\*|_)(.*?)\1")  # Matches *italic* or _italic_

    # Keep track of the position in text
    pos = 0
    for match in re.finditer(r"(\*\*\*|__|\*\*|\*|_)(.+?)\1", text):
        # Add text before the match
        if match.start() > pos:
            paragraph.add_run(text[pos : match.start()])

        style = match.group(1)  # Markdown style found (** or *)
        matched_text = match.group(2)  # Matched content

        # Apply formatting based on style
        run = paragraph.add_run(matched_text)
        if style in ("***", "___"):  # Bold italic
            run.bold = True
            run.italic = True
        elif style in ("**", "__"):  # Bold
            run.bold = True
        elif style in ("*", "_"):  # Italic
            run.italic = True

        # Update position
        pos = match.end()

    # Add any remaining text after the last match
    if pos < len(text):
        paragraph.add_run(text[pos:])


def create_word_table(comments, output_filename="revision_table.docx"):
    # Create a new Word document
    doc = Document()
    doc.add_heading("Revision Table", level=1)

    # Add table with three columns
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.autofit = False

    # Set column widths
    table.columns[0].width = Inches(1.0)
    table.columns[1].width = Inches(4.0)
    table.columns[2].width = Inches(4.0)

    # Add header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Comment-number"
    hdr_cells[1].text = "Comment"
    hdr_cells[2].text = "Response"

    # Populate table rows with formatted Markdown content
    for comment_number, comment_text, response_text in comments:
        row_cells = table.add_row().cells
        row_cells[0].text = comment_number

        # Add formatted Markdown text for comment and response
        comment_paragraph = row_cells[1].paragraphs[0]
        add_markdown_text(comment_paragraph, comment_text)

        response_paragraph = row_cells[2].paragraphs[0]
        add_markdown_text(response_paragraph, response_text)

    # Set page orientation to landscape
    section = doc.sections[0]
    new_width, new_height = section.page_height, section.page_width
    section.page_width = new_width
    section.page_height = new_height

    # Save the document
    doc.save(output_filename)
    print(f"Revision table saved to {output_filename}")


# Parse comments and create Word table
comments = parse_comments(lines)
create_word_table(comments)
